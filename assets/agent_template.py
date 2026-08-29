#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""סוכן התוכן — הריצה הלילית.

התבנית הזו מגיעה מסוכן שרץ בעסק חי ונשבר ב-14 דרכים שונות. כל הגנה כאן
מתקנת כשל אמיתי, ומי שעורך אותה כדאי שיקרא קודם את references/reliability.md.

שימוש:
    python3 agent.py           ריצה מלאה
    python3 agent.py --test    קובץ אחד, פלט אחד, בלי לעדכן מצב
"""
import json
import shutil
import re
import subprocess
import sys
import traceback
from datetime import datetime, timedelta
from pathlib import Path

HERE = Path(__file__).resolve().parent
CONFIG = json.loads((HERE / "config.json").read_text(encoding="utf-8"))

# מבנה התיקיות מועתק מסוכן שרץ בעסק חי חודשים, לא הומצא כאן.
# תוכן מסודר לפי סוג — "תביא לי קרוסלה" זה מקום אחד. התאריך בשם הקובץ.
# הכל בדרייב, כי שם בעל העסק כבר מחפש את הדברים שלו ומשם הוא מגיע
# אליהם גם מהנייד.
CONTENT_ROOT = Path(CONFIG["content_folder"])
FOLDERS = {
    "POV": CONTENT_ROOT / "POV",
    "ריל": CONTENT_ROOT / "רילס",
    "קרוסלה": CONTENT_ROOT / "קרוסלה",
    "פוסט": CONTENT_ROOT / "פוסט",
    "סטורי": CONTENT_ROOT / "סטורי",
}
DAILY_SUMMARY = CONTENT_ROOT / "_סיכום_יומי"
LEARNED = CONTENT_ROOT / "_מה_למדנו"  # מצטבר, לא נמחק — זה הנכס
MIRROR = CONTENT_ROOT / "_מראה"  # ניתוח השיחה עצמה. לבעל העסק, לא לתוכן.

# הארכיון: התמלול המלא נשמר, לא נמחק. גם אם יוחלף כלי התמלול מחר,
# ההיסטוריה נשארת ואפשר לחזור ולהוציא ממנה עוד.
ARCHIVE = Path(CONFIG.get("archive_folder") or (CONTENT_ROOT.parent / "פגישות לקוחות"))

MONTHS = ["ינואר","פברואר","מרץ","אפריל","מאי","יוני","יולי","אוגוסט","ספטמבר","אוקטובר","נובמבר","דצמבר"]


def month_dir(base, when=None):
    """base/2026/אוגוסט — תיקיית חודש נוצרת רק כשיש בה תוכן"""
    d = when or datetime.now()
    return base / str(d.year) / MONTHS[d.month - 1]


# ── הגשר לגיליון. אופציונלי: בלעדיו הסוכן עובד בדיוק אותו דבר ──────
PIPELINE = None
DRIVE_LINKS = {}   # נתיב מקומי -> לינק בדרייב
if CONFIG.get("pipeline_url"):
    try:
        import sys as _sys
        _sys.path.insert(0, str(HERE))
        from pipeline import Pipeline as _P
        PIPELINE = _P(CONFIG["pipeline_url"])
    except Exception as _e:
        PIPELINE = None

# המנוע מקומי בכוונה: state.json נכתב תוך כדי ריצה, וקובץ שהסנכרון
# תופס באמצע כתיבה הוא קובץ פגום.
LOGS = HERE / "logs"
STATE_FILE = HERE / "state.json"

TRANSCRIPTS = Path(CONFIG["transcripts_folder"])
VOICE_FILE = Path(CONFIG["voice_file"])

# דוגמאות טובות. כללים מלמדים מה אסור, דוגמאות מלמדות צורה.
# אופציונלי לגמרי — בלעדיו הסוכן עובד בדיוק אותו דבר, רק פחות טוב.
EXAMPLES_FILE = Path(CONFIG.get("examples_file") or (CONTENT_ROOT / "דוגמאות-טובות.md"))

# השאלות שחוזרות. נבנה מחדש בכל ריצה מכל הניתוחים — בלי state,
# כדי שטעות אחת לא תישאר לתמיד.
QUESTIONS_FILE = CONTENT_ROOT / "השאלות-שחוזרות.md"
SKILL_REFS = Path(CONFIG["skill_refs"])  # תיקיית references של הסקיל

CLAUDE = CONFIG.get("claude_path", "claude")
CLAUDE_TIMEOUT = 600
LOOKBACK_DAYS = CONFIG.get("lookback_days", 30)
# כמה פורמטים לכל פגישה. 3 פגישות ביום → 9 פיסות תוכן.
# כל פיסה מבוססת על פגישה אחת בלבד, בלי ערבוב — כשמערבבים שתי פגישות
# בפוסט אחד הזווית מתפשרת ואף אחת מהן לא נשמעת אמיתית.
PIECES_PER_MEETING = CONFIG.get("pieces_per_meeting", 3)

READABLE = {".txt", ".md", ".json", ".vtt", ".srt", ".docx"}

FORMATS = {
    # POV הוא פורמט בפני עצמו, לא תת-סוג של פוסט. יש לו חוקי קול
    # משלו (references/pov.md) והוא הפורמט שהכי מהר בונה חיבור.
    "POV": ["POV1", "POV2", "POV3"],
    "פוסט": [f"P{i}" for i in range(1, 11)],
    "ריל": [f"R{i}" for i in range(1, 9)],
    "קרוסלה": [f"C{i}" for i in range(1, 8)],
    "סטורי": [f"S{i}" for i in range(1, 5)],
}

# ── לוג ─────────────────────────────────────────────────────────────

LOG_FILE = None


def log(msg):
    line = f"[{datetime.now():%Y-%m-%d %H:%M:%S}] {msg}"
    print(line, flush=True)
    if LOG_FILE:
        try:
            with open(LOG_FILE, "a", encoding="utf-8") as f:
                f.write(line + "\n")
        except Exception:
            pass


def setup_logging():
    global LOG_FILE
    LOGS.mkdir(parents=True, exist_ok=True)
    LOG_FILE = LOGS / f"{datetime.now():%Y-%m-%d}.log"


def safe(fn, *a, **kw):
    """פעולה נלווית שכישלון בה לא צריך להפיל ריצה שלמה.

    למה: chmod ויצירת תיקיות נכשלים כשהסקריפט רץ מהמתזמן ולא מהטרמינל —
    למתזמן אין את אותן הרשאות. זה הפיל סוכן אמיתי 6 פעמים.
    """
    try:
        return fn(*a, **kw)
    except Exception as e:
        log(f"  (פעולה נלווית נכשלה, ממשיך: {e.__class__.__name__})")
        return None


# ── מצב ─────────────────────────────────────────────────────────────


def load_state():
    base = {"last_run": None, "processed_files": [], "format_history": {},
            "hooks_sent": {}, "produced": [], "produced_week": {}}
    if STATE_FILE.exists():
        try:
            s = json.loads(STATE_FILE.read_text(encoding="utf-8"))
            base.update(s)
        except Exception as e:
            log(f"state.json פגום ({e}) — מתחיל מחדש")
    return base


def save_state(state):
    STATE_FILE.write_text(
        json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8"
    )


# ── קלוד ────────────────────────────────────────────────────────────

JUNK_MARKERS = [
    "מופיע למעלה",
    "כפי שציינתי",
    "כבר השתמשתי",
    "as mentioned above",
    "I have everything I need",
    "Let me produce",
]


def looks_like_junk(text, min_len=800):
    """תופס פלט שהוא מטא-טקסט במקום תוצר.

    מודל מחזיר לפעמים "הניתוח מופיע למעלה" במקום הניתוח. בלי הבדיקה הזו
    זה נשמר, והלוג רושם ✅ על 300 בתים של כלום.
    """
    if not text or len(text.strip()) < min_len:
        return f"קצר מדי ({len(text.strip()) if text else 0} תווים)"
    opening = text.strip()[:400].lower()
    for m in JUNK_MARKERS:
        if m.lower() in opening:
            return f"נפתח במטא-טקסט: {m!r}"
    return None


def strip_preamble(text):
    """חותך מחשבות פנימיות שדלפו לפני התוכן.

    ראינו קופי שנשמר כשהוא פותח ב-"Good - the metaphor is verified".
    זו מחשבה של המודל, לא טקסט לפרסום.
    """
    lines = text.split("\n")
    start = 0
    for i, line in enumerate(lines[:6]):
        st = line.strip()
        if not st:
            continue
        if st.startswith(("#", "**", "##")):
            start = i
            break
        letters = [c for c in st if c.isalpha()]
        if letters and sum(1 for c in letters if "a" <= c.lower() <= "z") / len(letters) > 0.5:
            start = i + 1
    out = "\n".join(lines[start:]).strip(" -\n")
    return out or text.strip()


def ask_claude(prompt, label="claude", min_len=800):
    """שולח לקלוד, מנקה, מאמת, ומנסה שוב פעם אחת."""
    tmp = Path("/tmp/content_agent")
    tmp.mkdir(exist_ok=True)
    pf = tmp / f"{label}_{datetime.now():%Y%m%d%H%M%S%f}.txt"
    pf.write_text(prompt, encoding="utf-8")
    try:
        for attempt in (1, 2):
            # --allowedTools חיוני: בהרצה אוטומטית אין מי שיאשר הרשאה
            # אינטראקטיבית, וקלוד יחזיר תשובה מנומסת עם קוד יציאה 0.
            r = subprocess.run(
                [CLAUDE, "-p", f"קרא את {pf} ובצע. החזר רק את הפלט הסופי.",
                 "--allowedTools", "Read,Glob,Grep"],
                capture_output=True, text=True, timeout=CLAUDE_TIMEOUT,
            )
            if r.returncode != 0:
                log(f"  קלוד נכשל (ניסיון {attempt}): {r.stderr[:200]}")
                continue
            out = strip_preamble(r.stdout.strip())
            bad = looks_like_junk(out, min_len)
            if not bad:
                return out
            log(f"  פלט לא תקין ({bad}) — ניסיון {attempt}")
        return None
    finally:
        safe(pf.unlink)


# ── סודיות ──────────────────────────────────────────────────────────


def anonymize(text, known_names=None):
    """מחליף שמות בתוויות ניטרליות, לפני שהתמלול נשלח לניתוח.

    מה שלא נכנס לא יכול לצאת. זו ההגנה הזולה והאמינה ביותר.

    ⚠️ שתי שכבות, ושתיהן נדרשות:

    1. **תוויות דובר** בתחילת שורה ("שם: טקסט").
    2. **שמות בתוך הדיבור עצמו** — "ודניאל בסוף יוני מתפנה". זו הדליפה
       שנתפסה בבדיקה: שם של אדם שלישי הוזכר עשר פעמים בגוף השיחה, עבר
       את הניקוי, והגיע עד למייל של בעל העסק.

    ⚠️ התוויות ניטרליות בכוונה — [דובר 1], לא [בעל העסק].
    ניחוש לפי סדר הדיבור נכשל: בשיחות רבות הלקוח פותח. ניחוש שגוי כאן
    מייחס את הכאבים של הלקוח לבעל העסק, וכל התוכן יוצא הפוך.
    """
    speakers = {}
    pat = re.compile(r"^(?:\[[^\]]+\]\s*)?([A-Za-z֐-׿][\w֐-׿' ]{1,30}):\s")
    for line in text.split("\n"):
        m = pat.match(line.strip())
        if m:
            speakers.setdefault(m.group(1).strip(), None)
    for i, name in enumerate(speakers, 1):
        speakers[name] = f"[דובר {i}]"
    for name, tag in speakers.items():
        text = re.sub(rf"(?<![\w֐-׿]){re.escape(name)}(?![\w֐-׿])", tag, text)

    # שכבה שנייה: כל שם ידוע, בכל מקום בטקסט — כולל חלקי שם.
    # "Daniel Shenhar" צריך לתפוס גם "דניאל" וגם "Daniel" לחוד.
    for full in sorted(known_names or [], key=len, reverse=True):
        parts = [full] + [w for w in full.split() if len(w) > 2]
        for w in sorted(set(parts), key=len, reverse=True):
            text = re.sub(rf"(?<![\w֐-׿]){re.escape(w)}(?![\w֐-׿])", "[הושמט]", text)

    return text


# תקרת התמלול. **לא 40,000.**
# בבדיקה על 173 תמלולים אמיתיים: החציון 106,779 תווים, המקסימום 230,822,
# ו-159 מתוך 173 מעל 40K. כלומר בתקרה הישנה הסוכן ניתח שליש מכל פגישה,
# **וסוף הפגישה — שם נסגרות ההחלטות — מעולם לא נקרא.**
TRANSCRIPT_LIMIT = 180000


def fit_transcript(text, limit=TRANSCRIPT_LIMIT):
    """מחזיר (טקסט, הערה). כשחייבים לקצץ — לוקחים משלושת החלקים.

    חיתוך מההתחלה בלבד הוא הגרוע ביותר: הפתיחה היא חימום, האמצע הוא
    התוכן, והסוף הוא ההחלטות והסיכום. אלה בדיוק החלקים שהתוכן צריך.
    """
    n = len(text)
    if n <= limit:
        return text, ""
    third = limit // 3
    head = text[:third]
    mid_start = (n - third) // 2
    mid = text[mid_start:mid_start + third]
    tail = text[-third:]
    note = (f"⚠️ הפגישה ארוכה ({n:,} תווים). כאן ההתחלה, האמצע והסוף — "
            "לא ברצף. אל תניח שהחסר לא היה.")
    return (f"{head}\n\n[...קטע הושמט...]\n\n{mid}"
            f"\n\n[...קטע הושמט...]\n\n{tail}"), note


def meeting_date(f):
    """התאריך של הפגישה, לא של הקריאה.

    למה זה חשוב: השאלות החוזרות נספרות לפי תאריך. אם כל הקבצים
    מתויגים ביום שבו הסוכן קרא אותם, הם נראים לו כפגישה אחת —
    **ואז שום שאלה לא יכולה לחזור.** זה נתפס בהרצה אמיתית על
    שלוש פגישות: "17 שאלות מ-1 פגישות · 0 שאלות חוזרות".

    זה קורה בכל הרצה ראשונה על ארכיון קיים, וזה בדיוק מה שתלמיד
    עושה ביום הראשון.

    כמעט כל כלי תמלול שם את התאריך בשם הקובץ. משתמשים בו כשהוא
    שם, ונופלים ל-mtime רק כשאין.
    """
    m = re.search(r"(20\d{2})[-_.]?(\d{2})[-_.]?(\d{2})", f.name)
    if m:
        y, mo, d = m.groups()
        if "01" <= mo <= "12" and "01" <= d <= "31":
            return f"{y}-{mo}-{d}"
    return datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d")


def has_leaked_names(text):
    """שכבה שנייה: שואל את קלוד אם נשאר שם מזהה בפלט.

    למה לא regex: רשימת שמות קבועה תופסת עשרה שמות ומפספסת אלפים, וגם
    לא תתפוס זיהוי עקיף ("הבעלים של הרשת הגדולה בצפון"). הסודיות היא
    הדבר הקריטי ביותר כאן, וזה לא מקום לחסוך בו קריאה אחת.

    מחזיר את מה שנמצא, או None אם נקי.
    """
    check = ask_claude(
        f"""האם בטקסט הבא מופיע שם של אדם, שם של עסק, שם מקום, או פרט
שמאפשר לזהות לקוח מסוים?

ציטוט בלי שם הוא תקין. תיאור כללי ("לקוחה בתחום היופי") תקין.
מה שלא תקין: שם פרטי, שם משפחה, שם עסק, עיר, או צירוף שמזהה אדם ספציפי.

ענה במילה אחת "נקי", או בשם/הפרט שמצאת. בלי הסבר.

---
{text[:6000]}""",
        label="privacy_check",
        min_len=0,
    )
    if not check:
        # אם הבדיקה עצמה נכשלה — לא מניחים שנקי
        log("  ⚠️ בדיקת הסודיות לא רצה — מתייחס כאילו יש חשד")
        return "בדיקה נכשלה"
    # מנקים עיצוב לפני ההשוואה: קלוד עונה לפעמים "**נקי**" עם הדגשה,
    # ובדיקה על המחרוזת הגולמית פוסלת אז טקסט תקין לגמרי.
    # זו מלכודת 6 (הבודק שפוסל עבודה טובה) במסווה חדש.
    ans = re.sub(r"[*_`#\s]+", " ", check).strip()
    return None if ans.startswith("נקי") else ans[:60]


def upload_to_drive(local_path, kind, drive_folder_name):
    """מעלה את הקובץ לגוגל דרייב דרך קלוד, בלי אפליקציית דרייב על המחשב.

    למה לא תיקייה מסונכרנת: זה תלוי באפליקציה שרצה. אצל בעל עסק אמיתי
    האפליקציה נמחקה (תפסה מקום בדיסק), והתוצאה הייתה שחודשיים של תוכן
    ישבו על המחשב ולא הגיעו לענן — **בלי שום שגיאה.** הסוכן דיווח
    "נשמר", והוא באמת נשמר. הוא פשוט לא הגיע ליעד.

    ככה זה לא תלוי בכלום חוץ מהחיבור שכבר קיים בקלוד קוד.
    """
    if not CONFIG.get("upload_to_drive", True):
        return True
    content = local_path.read_text(encoding="utf-8")
    root = CONFIG.get("drive_folder_name") or CONTENT_ROOT.name
    root_id = CONFIG.get("drive_folder_id", "").strip()

    # ⚠️ המסלול המלא חייב להיות מפורש, ועדיף מזהה תיקייה.
    # בבדיקה אמיתית הפרומפט אמר רק "העלה לתיקייה 'רילס'", וקלוד מצא
    # תיקייה בשם הזה במקום אחר בדרייב והעלה לשם — לתוך התוכן האמיתי
    # של מישהו אחר. שם תיקייה לבדו הוא לא כתובת.
    where = (
        f"לתיקייה עם המזהה {root_id} (זו תיקיית '{root}'), במסלול '{drive_folder_name}'"
        if root_id
        else f"למסלול '{drive_folder_name}' שנמצא **בתוך** התיקייה '{root}'"
    )
    prompt = (
        f"העלה לגוגל דרייב קובץ בשם '{local_path.stem}' {where}.\n\n"
        f"⚠️ אם המסלול '{drive_folder_name}' לא קיים בתוך '{root}' — צור את כל התיקיות בדרך.\n"
        f"⚠️ בסוף החזר את הקישור לקובץ שנוצר, בשורה נפרדת.\n"
        f"⚠️ אל תעלה לתיקייה בשם דומה שנמצאת במקום אחר בדרייב. "
        f"רק בתוך '{root}'.\n\n"
        f"התוכן:\n\n{content}"
    )
    try:
        r = subprocess.run(
            [CLAUDE, "-p", prompt, "--allowedTools",
             "mcp__claude_ai_Google_Drive__create_file,mcp__claude_ai_Google_Drive__search_files"],
            capture_output=True, text=True, timeout=300,
        )
    except Exception as e:
        log(f"  ⚠️ ההעלאה לדרייב קרסה: {e.__class__.__name__}")
        return False

    ans = (r.stdout or "").lower()
    if r.returncode != 0 or any(s in ans for s in
                                ["לא הצלחתי", "אין לי גישה", "צריך לאשר", "failed", "permission"]):
        log(f"  ⚠️ ההעלאה לדרייב נכשלה. הקובץ נשמר מקומית: {local_path.name}")
        return False
    m = re.search(r"https://(?:docs|drive)\.google\.com/\S+", r.stdout or "")
    if m:
        DRIVE_LINKS[str(local_path)] = m.group(0).rstrip(").,\"'")
        log("  ☁️ הועלה לדרייב")
    else:
        # הועלה, אבל בלי קישור. הגיליון יקבל את הנתיב המקומי
        log("  ☁️ הועלה לדרייב (בלי קישור בתשובה)")
    return True


def save_docx(md_text, out_path, title=None):
    """שומר גם גרסת Word, עם יישור לימין.

    למה: בעל עסק לא טכני עורך ב-Word או ב-Pages, לא בעורך טקסט. הסוכן
    שרץ בעסק חי מייצר את שתי הגרסאות, וזו שנפתחת בפועל היא ה-DOCX.

    נכשל בשקט אם הספרייה לא מותקנת — זה נחמד־שיהיה, לא קריטי, ואסור
    שיפיל ריצה שלמה.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        return False
    try:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.element.rPr.rFonts.set(qn("w:cs"), "Arial")
        if title:
            h = doc.add_heading(title, level=1)
            h.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        for raw in md_text.split("\n"):
            line = raw.rstrip()
            if line.startswith("#"):
                lvl = min(len(line) - len(line.lstrip("#")), 4)
                para = doc.add_heading(line.lstrip("# ").strip(), level=lvl)
            else:
                para = doc.add_paragraph(line.lstrip("- ") if line.startswith("- ") else line)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            para.paragraph_format.right_to_left = True
        doc.save(str(out_path))
        return True
    except Exception as e:
        log(f"  (DOCX נכשל, ה-MD נשמר: {e.__class__.__name__})")
        return False


# ── קריאת קבצים ─────────────────────────────────────────────────────


def read_transcript(path):
    """מוציא טקסט מכל פורמט שכלי תמלול מייצרים.

    מחזיר (טקסט, שמות_ידועים). השמות נאספים מרשימת הדוברים — בקובץ JSON
    הם יושבים בשדה נפרד, ואם לא אוספים אותם משם הם נכנסים לטקסט כמו שהם.
    """
    names = set()
    try:
        if path.suffix == ".json":
            d = json.loads(path.read_text(encoding="utf-8"))
            tr = d.get("transcript", d)
            for s in tr.get("speakers", []):
                n = (s.get("name") or "").strip()
                if n:
                    names.add(n)
            segs = tr.get("segments") or d.get("segments") or []
            if segs:
                text = "\n".join(
                    f"{s.get('speaker_id', s.get('speaker', '?'))}: {s.get('text', '')}"
                    for s in segs
                )
                return text, names
            return json.dumps(d, ensure_ascii=False), names

        if path.suffix == ".docx":
            # python-docx קודם: הוא עובד על כל מערכת הפעלה.
            # textutil קיים רק במק — תלמיד על Windows היה מקבל
            # מחרוזת ריקה, והקובץ היה מדולג כ"קצר מדי". בשקט.
            try:
                from docx import Document
                doc = Document(str(path))
                out = "\n".join(par.text for par in doc.paragraphs)
                for tbl in doc.tables:
                    for row in tbl.rows:
                        out += "\n" + " | ".join(c.text for c in row.cells)
                if out.strip():
                    return out, names
            except Exception:
                pass
            if shutil.which("textutil"):
                r = subprocess.run(
                    ["textutil", "-convert", "txt", "-stdout", str(path)],
                    capture_output=True, text=True, timeout=60)
                return r.stdout, names
            log(f"  ⚠️ אי אפשר לקרוא {path.name} — "
                "להתקין python-docx (pip install python-docx)")
            return "", names

        if path.suffix in {".vtt", ".srt"}:
            raw = path.read_text(encoding="utf-8", errors="ignore")
            keep = [l for l in raw.split("\n")
                    if l.strip() and "-->" not in l and not l.strip().isdigit()
                    and l.strip() != "WEBVTT"]
            return "\n".join(keep), names

        return path.read_text(encoding="utf-8", errors="ignore"), names
    except Exception as e:
        log(f"  לא הצלחתי לקרוא {path.name}: {e}")
        return None, names


def find_new(state):
    """חלון קבוע אחורה, וסינון לפי מה שכבר טופל.

    לא "מאז הריצה האחרונה": תמלול שהושלם באיחור נעלם אז לתמיד.
    """
    if not TRANSCRIPTS.exists():
        raise RuntimeError(f"תיקיית התמלולים לא קיימת: {TRANSCRIPTS}")
    cutoff = (datetime.now() - timedelta(days=LOOKBACK_DAYS)).timestamp()
    done = set(state["processed_files"])
    out = []
    for f in TRANSCRIPTS.rglob("*"):
        if f.is_file() and f.suffix.lower() in READABLE and f.name not in done:
            if f.stat().st_mtime >= cutoff:
                out.append(f)
    return sorted(out, key=lambda p: p.stat().st_mtime)


# ── רוטציה ──────────────────────────────────────────────────────────


def pick_audience_layer(kind, used_today, layers):
    """כל פיסת תוכן מכוונת לפלח קהל אחר, ומתחלפת ביניהם.

    למה זה חשוב: בלי זה כל התוכן של אותו יום מדבר לאותו אדם, והפיד
    נשמע כמו תקליט שנתקע. הסוכן שרץ בעסק חי עושה את זה מהיום הראשון,
    ובלי זה התוכן נראה חוזר על עצמו גם כשהפורמטים מתחלפים.

    הפלחים מגיעים מקובץ הקול (החלק "למי אני מוכר"). אם אין שם פילוח,
    מחזירים None והתוכן פשוט לא מכוון לפלח מסוים.
    """
    if not layers:
        return None
    # לכל סוג תוכן סדר עדיפויות אחר, כדי שיום אחד יכסה כמה פלחים
    offset = {"POV": 0, "ריל": 1, "קרוסלה": 2, "פוסט": 3, "סטורי": 4}.get(kind, 0)
    order = layers[offset % len(layers):] + layers[: offset % len(layers)]
    for l in order:
        if l not in used_today:
            return l
    return order[0]


def extract_layers(voice_text):
    """שולף את פלחי הקהל מקובץ הקול, אם יש שם פילוח.

    מחפש כותרת "למי אני מוכר" ולוקח ממנה שורות רשימה. זה לא מדע מדויק,
    ובכוונה סלחני: עדיף להחזיר רשימה ריקה מאשר לנחש פלחים שלא קיימים.
    """
    m = re.search(r"##\s*למי אני מוכר(.*?)(?=\n##|\Z)", voice_text, re.S)
    if not m:
        return []
    out = []
    for line in m.group(1).split("\n"):
        s = line.strip(" -*•\t")
        if 12 < len(s) < 90 and not s.startswith("#"):
            out.append(s[:60])
    return out[:4]


def pick_format(kind, history):
    """הפורמט שהכי מזמן לא היה. לא 'הראשון הפנוי' — זה נותן סדר קבוע."""
    avail = FORMATS[kind]
    never = [f for f in avail if f not in history]
    if never:
        return never[0]
    return sorted(avail, key=lambda f: history.get(f, ""))[0]


# ── הריצה ───────────────────────────────────────────────────────────


def run(test_mode=False):
    setup_logging()
    log("=" * 60)
    log("ריצה התחילה" + (" (מצב בדיקה)" if test_mode else ""))

    for d in [DAILY_SUMMARY, LEARNED, ARCHIVE, MIRROR]:
        safe(d.mkdir, parents=True, exist_ok=True)

    if not VOICE_FILE.exists():
        raise RuntimeError(f"קובץ הקול לא נמצא: {VOICE_FILE}")
    voice = VOICE_FILE.read_text(encoding="utf-8")

    examples = ""
    if EXAMPLES_FILE.exists():
        examples = EXAMPLES_FILE.read_text(encoding="utf-8").strip()
        n = examples.count("---")
        log(f"📎 דוגמאות טובות: {EXAMPLES_FILE.name} ({n} דוגמאות)")
    else:
        log("ℹ️ אין קובץ דוגמאות. הכתיבה תסתמך על הכללים בלבד")

    state = load_state()
    new_files = find_new(state)
    if test_mode:
        new_files = new_files[:1]
    log(f"נמצאו {len(new_files)} קבצים חדשים")

    if not new_files:
        log("אין חומר חדש")
        safe(DAILY_SUMMARY.mkdir, parents=True, exist_ok=True)
        (DAILY_SUMMARY / f"{datetime.now():%Y-%m-%d}.md").write_text(
            f"# {datetime.now():%Y-%m-%d}\n\nלא היו פגישות חדשות", encoding="utf-8")
        return

    analyses, mirrors = [], []
    for i, f in enumerate(new_files, 1):
        log(f"[{i}/{len(new_files)}] {f.name}")
        raw, known = read_transcript(f)
        if not raw or len(raw) < 200:
            log("  קצר מדי או לא נקרא, מדלג")
            continue

        clean = anonymize(raw, known)
        prompt = build_analysis_prompt(clean, voice)
        result = ask_claude(prompt, "analysis", min_len=800)
        if not result:
            log("  הניתוח נכשל — לא מסמן כמטופל, ינוסה שוב מחר")
            continue

        # הניתוח נשמר ומוזן למייל היומי, אז הוא חייב לעבור את אותה
        # בדיקה כמו הקופי. בלי זה שם שהוסק מהקשר מגיע ישר לתיבה.
        leaked = has_leaked_names(result)
        if leaked:
            log(f"  ⚠️ נמצא שם בניתוח ({leaked}) — מנקה")
            result = ask_claude(
                f"""הסר מהטקסט הבא כל שם של אדם, עסק או מקום, והחלף
ב-[הושמט]. אל תשנה שום דבר אחר. החזר את הטקסט המתוקן בלבד.

---
{result}""", "analysis_scrub", min_len=500) or result
            if has_leaked_names(result):
                log("  ❌ הניתוח עדיין מכיל שם — מדלג על הפגישה")
                continue

        # ── ניתוח המראה ────────────────────────────────────────────
        # רץ אחרי הניתוח הרגיל, ולא במקומו. כישלון כאן לא מפיל כלום —
        # התוכן כבר נשמר.
        stamp = meeting_date(f)
        mp = build_mirror_prompt(raw)
        if mp:
            mirror = ask_claude(mp, "mirror", min_len=1500)
            if mirror:
                safe(MIRROR.mkdir, parents=True, exist_ok=True)
                safe((MIRROR / f"{stamp}__{f.stem[:40]}.md").write_text,
                     mirror, encoding="utf-8")
                log("  🪞 ניתוח מראה נשמר")
            else:
                log("  (ניתוח המראה לא הצליח — ממשיך)")
                mirror = ""
        else:
            mirror = ""

        out = LEARNED / f"{stamp}__{f.stem[:40]}__{abs(hash(f.name)) % 100000}.md"
        out.write_text(result, encoding="utf-8")
        # שתי הרשימות גדלות **יחד ובשורות צמודות**, בכוונה.
        # אם הן יוצאות מסנכרון, המראה של פגישה אחת מתחברת לניתוח של
        # אחרת — והתוכן ייצא מדויק בטון ושגוי בעובדות. בלי שגיאה.
        analyses.append(result)
        mirrors.append(mirror)
        if not test_mode:
            state["processed_files"].append(f.name)
            # התמלול המלא עובר לארכיון ולא נמחק. זה מה שמאפשר לחזור
            # לשיחה ישנה ולהוציא ממנה עוד, גם אחרי החלפת כלי תמלול.
            dest_dir = month_dir(ARCHIVE, datetime.fromtimestamp(f.stat().st_mtime))
            safe(dest_dir.mkdir, parents=True, exist_ok=True)
            dest = dest_dir / f.name
            if dest.exists():
                dest = dest_dir / f"{f.stem}__{abs(hash(f.name)) % 10000}{f.suffix}"
            if safe(shutil.move, str(f), str(dest)) is None:
                log("  ⚠️ ההעברה לארכיון נכשלה — הקובץ נשאר במקומו")
            else:
                log(f"  📦 לארכיון: {dest.parent.name}/{dest.name}")
        log(f"  ✅ נשמר ({len(result):,} תווים)")

    if not analyses:
        raise RuntimeError("אף קובץ לא עובד בהצלחה")

    ALL_KINDS = ["POV", "ריל", "קרוסלה", "פוסט", "סטורי"]
    layers = extract_layers(voice)
    if layers:
        log(f"פלחי קהל מקובץ הקול: {len(layers)}")
    used_layers = []
    today = f"{datetime.now():%Y-%m-%d}"
    created = []
    hooks = []      # מה שיירשם בגיליון בסוף הריצה

    # כל פגישה מקבלת את מלוא תשומת הלב שלה
    jobs = []
    for m_idx in range(len(analyses)):
        for k in range(1 if test_mode else PIECES_PER_MEETING):
            jobs.append((ALL_KINDS[(m_idx + k) % len(ALL_KINDS)], m_idx))
        if test_mode:
            break
    assert len(mirrors) == len(analyses), "המראות והניתוחים יצאו מסנכרון"
    log(f"{len(analyses)} פגישות → {len(jobs)} פיסות תוכן")

    # ── השאלות שחוזרות ─────────────────────────────────────────────
    # רץ אחרי הניתוחים ולפני הכתיבה, כדי שהכתיבה תצא משאלה אמיתית
    # ולא מזווית שהומצאה. כישלון כאן לא מפיל ריצה — נופלים למסרים
    # שבקובץ הקול, בדיוק כמו בחודש הראשון.
    questions_text = ""
    try:
        from questions import build as build_questions
        log("מחפש שאלות שחוזרות")
        n, qpath = build_questions(LEARNED, QUESTIONS_FILE, CLAUDE, log)
        if qpath and qpath.exists():
            questions_text = qpath.read_text(encoding="utf-8")
    except Exception as e:
        log(f"  ⚠️ חיפוש השאלות נכשל, ממשיך בלעדיו: {e.__class__.__name__}")

    for kind, m_idx in jobs:
        code = pick_format(kind, state["format_history"])
        layer = pick_audience_layer(kind, used_layers, layers)
        log(f"כותב {kind} ({code}) מפגישה {m_idx + 1}" + (f" · {layer[:30]}" if layer else ""))
        prompt = build_content_prompt(
            kind, code, analyses[m_idx], voice, questions_text, layer,
            examples, mirrors[m_idx] if m_idx < len(mirrors) else ""
        )
        content = ask_claude(prompt, f"content_{kind}", min_len=300)
        if not content:
            log(f"  {kind} נכשל, מדלג")
            continue

        leaked = has_leaked_names(content)
        if leaked:
            log(f"  ⚠️ נמצא שם בפלט ({leaked}) — כותב מחדש")
            content = ask_claude(
                prompt + "\n\n⚠️ אסור בהחלט להזכיר שמות של אנשים.",
                f"content_{kind}_retry", min_len=300)
            if not content or has_leaked_names(content):
                log(f"  ❌ {kind} עדיין מכיל שם — מדלג")
                continue

        folder = month_dir(FOLDERS[kind])
        safe(folder.mkdir, parents=True, exist_ok=True)
        p = folder / f"{today}_{code}_{m_idx + 1}.md"
        p.write_text(content, encoding="utf-8")
        save_docx(content, p.with_suffix(".docx"), title=f"{kind} · {today}")
        upload_to_drive(p, kind, f"{FOLDERS[kind].name}/{folder.parent.name}/{folder.name}")
        created.append(p)
        if layer:
            used_layers.append(layer)
            if len(used_layers) >= len(layers):
                used_layers = []
        if not test_mode:
            state["format_history"][code] = today
        hooks.append({"kind": kind, "path": p, "text": content})
        log(f"  ✅ נשמר ב{FOLDERS[kind].name}/{folder.parent.name}/{folder.name}/")

    # סיכום יומי — מה נוצר היום, במקום אחד
    if created:
        safe(DAILY_SUMMARY.mkdir, parents=True, exist_ok=True)
        lines = [f"# {today}", "", f"נותחו {len(analyses)} פגישות", "", "## נוצר היום", ""]
        lines += [f"- {p.parent.name}/{p.name}" for p in created]
        (DAILY_SUMMARY / f"{today}.md").write_text("\n".join(lines), encoding="utf-8")

    # ── שורה בגיליון לכל פיסת תוכן ──────────────────────────────────
    if hooks and PIPELINE:
        try:
            items = [{
                "תאריך": today,
                "סוג": h["kind"],
                "הוק": first_line(h["text"]),
                "לינק": drive_link(h["path"]),
            } for h in hooks]
            res = PIPELINE.add(items)
            # ההוק המקורי נשמר לפי הלינק (מפתח ייחודי לכל פיסה).
            # הריצה השבועית משווה מולו כדי לזהות מה בעל העסק שכתב.
            if not test_mode:
                sent = state.setdefault("hooks_sent", {})
                for it in items:
                    if it["לינק"]:
                        sent[it["לינק"]] = it["הוק"]
            log(f"📋 {res.get('added', 0)} שורות נכנסו לגיליון ({res.get('tab', '')})")
            if not test_mode:
                state["pipeline_rows"] = state.get("pipeline_rows", 0) + res.get("added", 0)
        except Exception as e:
            # הגיליון נכשל, התוכן כבר שמור. לא מפילים ריצה בגלל זה
            log(f"⚠️ הכתיבה לגיליון נכשלה: {e}")
            log("   התוכן שמור בדרייב. אפשר לתקן את הגשר ולהריץ --sync")
    elif hooks and not PIPELINE:
        log("ℹ️ אין גשר לגיליון. התוכן נשמר בדרייב בלבד")

    # ── שער האישור ─────────────────────────────────────────────────
    # מה שבעל העסק סימן בגיליון עובר להפקה. רק מה שסימן.
    # ברירת המחדל היא תמיד אי-הוצאה: לא סימן = לא נוצר = לא עלה.
    if PIPELINE and not test_mode:
        try:
            import gate
            state.setdefault("weekly_cap",
                             CONFIG.get("weekly_production_cap", gate.DEFAULT_CAP))
            log("שער האישור — בודק מה סומן")
            gate.run(PIPELINE, state, CONFIG.get("producer_cmd"), log)
        except Exception as e:
            log(f"⚠️ שער האישור נכשל, התוכן שמור: {e.__class__.__name__}")

    if not test_mode:
        state["last_run"] = datetime.now().isoformat()
        save_state(state)

    log("=" * 60)
    log(f"סיום. נוצרו {len(created)} תכנים ב{CONTENT_ROOT.name}/")


def first_line(text, limit=90):
    """ההוק שנרשם בגיליון — המשפט שבאמת פותח, לא הכותרת הטכנית.

    קבצי התוכן נפתחים בהערת פורמט ובכותרת פנימית ("קרוסלה C2 — ..."),
    ואם לוקחים את השורה הראשונה מקבלים בגיליון שורה שלא אומרת כלום
    לבעל העסק. לכן קודם מחפשים שדה מפורש, ורק אחר כך שורת תוכן.
    """
    LABELS = ("כותרת:", "הוק:", "פתיח:", "שקופית 1", "שורה 1")
    for raw in text.splitlines():
        line = raw.strip().strip("*_").strip()
        for lab in LABELS:
            if line.startswith(lab) or line.startswith("**" + lab):
                val = line.split(lab, 1)[1].strip(" *:—-")
                if len(val) >= 8:
                    return val[:limit].rstrip(" ,.:;-")
    for raw in text.splitlines():
        line = raw.strip()
        if not line or line.startswith(("<!--", "#", "---", "===", "|")):
            continue
        line = line.strip("*_").strip()
        if len(line) >= 12:
            return line[:limit].rstrip(" ,.:;-")
    return text.strip()[:limit]


def drive_link(path):
    """לינק לקובץ, אם הועלה לדרייב. אחרת הנתיב המקומי"""
    return DRIVE_LINKS.get(str(path), str(path))


def build_analysis_prompt(transcript, voice):
    _body, _note = fit_transcript(transcript)
    return f"""אתה מנתח פגישת לקוח עבור בעל עסק, כדי ללמוד על הקהל שלו.

קרא קודם: {SKILL_REFS}/privacy.md

⚠️ הכלל הקריטי: **דפוסים וציטוטים כן, שמות ופרטים מזהים לא.**
ציטוט מותר ואפילו רצוי — מה שאסור זה לייחס אותו לאדם מזוהה.

הקול והעסק של בעל העסק:
{voice[:3000]}

הדוברים מסומנים [דובר 1], [דובר 2] וכו'. **קודם כל זהה מההקשר מי מהם
בעל העסק ומי הלקוח** — בעל העסק הוא זה שמייעץ, שואל שאלות אבחון, ומציע
פתרונות. אל תניח שהדובר הראשון הוא בעל העסק; בשיחות רבות הלקוח פותח.

**שני הצדדים של השיחה הם חומר גלם, ולא אותו סוג של חומר:**
הלקוח נותן את הכאב והשאלה. בעל העסק נותן את האבחון, התשובה והסיפור.
תוכן שבנוי רק על הצד של הלקוח מתאר בעיה ולא מראה מי יודע לפתור אותה.

החזר בדיוק במבנה הזה:

## על מה דיברו

## הכאבים שעלו
במילים של הלקוח

## השאלות ששאלו

## ציטוטי זהב של הלקוח (5-10, verbatim, בלי מקור)
הכאב, הפחד, ההתנגדות — כפי שנאמרו

## ציטוטי זהב של בעל העסק (5-10, verbatim)
האבחון, התובנה, ההוראה, ה"רגע שבו נפל האסימון". **זה החומר שהופך
תוכן מתיאור בעיה לעמדה של מי שיודע.**

## ⭐ סיפורים ודוגמאות אישיות שבעל העסק סיפר
כל רגע שהוא סיפר על עצמו: מהעבר שלו, החלטה שקיבל, טעות שעשה, או
"גם אני הייתי שם". **חלץ אותם verbatim, זה הזהב לחיבור עם הקהל.**
אם לא היו בפגישה הזו — כתוב "לא היו".

## הבעיות שבעל העסק פותר
מה הוא אבחן, ואיזה פתרון או דרך חשיבה הציע

## איך הם מנסחים את זה — מילים וביטויים
השפה המדויקת של הלקוח: סלנג, ביטויים חוזרים, איך הוא קורא לדברים

## 4 זוויות אפשריות לתוכן
לכל אחת, ציין באיזה ציטוט ובאיזה סיפור אפשר להשתמש

---
התמלול: {_note}

{_body}
"""


def build_mirror_prompt(transcript):
    """ניתוח מראה — מה באמת קרה בשיחה.

    **זה לא פרומפט התוכן, וזה בכוונה.** פרומפט הניתוח מוציא חומר גלם
    לכתיבה: ציטוטים, כאבים, שאלות. זה מוציא משהו אחר לגמרי — הדינמיקה,
    מה שבר את הזרימה, ומה שהיה צריך להיאמר ולא נאמר.

    הוא שווה לבעל העסק **גם אם לא יפרסם מילה**, וזה מה שהופך את
    הסוכן ממכונת תוכן לכלי עבודה.
    """
    body, note = fit_transcript(transcript)
    spec = ""
    try:
        spec = (SKILL_REFS / "mirror.md").read_text(encoding="utf-8")
    except Exception:
        return None
    return f"""{spec}

---
התמלול: {note}

{body}
"""


def build_content_prompt(kind, code, analysis, voice, questions="", layer=None, examples="", mirror=""):
    format_ref = (
        f"{SKILL_REFS}/pov.md — הפורמט הזה הוא POV. קרא את הקובץ במלואו,\n"
        "   יש לו חוקי קול משלו ונקודת כשל שקל ליפול בה"
        if kind == "POV"
        else f"{SKILL_REFS}/formats.md — מצא את הפרומפט {code} והפעל אותו בדיוק"
    )
    mirror_block = (
        "\n" + "=" * 60 + "\n"
        "🪞 ניתוח השיחה — מה באמת קרה שם:\n\n"
        f"{mirror}\n"
        + "=" * 60 + "\n\n"
        "**מכאן מגיע החומר הכי חזק.** בניתוח הזה יש הדינמיקה, הפחדים,\n"
        "מה שבר את הזרימה, ומה שהיה צריך להיאמר ולא נאמר.\n"
        "השורות האלה נוגעות יותר מכל תיאור של בעיה.\n"
        if mirror.strip() else ""
    )
    questions_block = (
        "=" * 60 + "\n"
        "⭐ השאלות שלקוחות שואלים אותו שוב ושוב:\n\n"
        f"{questions}\n"
        + "=" * 60 + "\n\n"
        "⚠️⚠️ **זו נקודת הפתיחה של הפיסה הזו:**\n\n"
        "**ההוק הוא שאלה חוזרת — בניסוח של הלקוח, לא בניסוח שלך.**\n"
        "שאלה שנשאלה חמש פעמים כבר הוכיחה שהיא מטרידה אנשים אמיתיים.\n"
        "היא לא צריכה שיפור, היא צריכה שיישמעו אותה.\n\n"
        "**הגוף הוא התשובה של בעל העסק מהניתוח.** מה שהוא ענה בפועל,\n"
        "בפגישה, במילים שלו. לא הסבר כללי על הנושא.\n\n"
        "בחר את השאלה שהכי מתאימה לניתוח שלפניך. אם אף שאלה לא מתחברת\n"
        "לחומר של הפגישה הזו — עדיף לצאת מהחומר, ולא למתוח שאלה בכוח.\n"
        if questions.strip() else
        "**עדיין אין מספיק פגישות כדי לדעת מה חוזר.**\n"
        "צא מ-`## המסרים שלי` שבקובץ הקול — מה שהוא הכי רוצה שיבינו.\n"
    )
    examples_block = (
        "\n" + "=" * 60 + "\n"
        "דוגמאות שבעל העסק בחר בעצמו — כאלה שהוא חושב שהן מעולות:\n\n"
        f"{examples}\n"
        + "=" * 60 + "\n\n"
        "⚠️⚠️ **מה ללמוד מהן ומה לא — זה קריטי:**\n\n"
        "✅ **קח מהן צורה בלבד:** אורך השורה · הקצב · איפה נשבר המשפט ·\n"
        "   כמה מוקדם מגיע המספר · איך נסגר הקפשן\n\n"
        "⛔ **אל תיקח מהן מילים, סלנג, סיפורים, מספרים או זהות.**\n"
        "   כל אלה באים **רק** מקובץ הקול ומהניתוח של הפגישה.\n\n"
        "אם דוגמה כתובה בסגנון של מישהו אחר — הצורה שלה מותרת, המילים שלה אסורות.\n"
        if examples else ""
    )
    layer_line = (
        f"\n⚠️ הפיסה הזו מכוונת לפלח קהל מסוים: **{layer}**\n"
        "כתוב אליו, לא לכולם. פיסה שמנסה לדבר לכל הקהל לא מדברת לאף אחד.\n"
        if layer else ""
    )
    return f"""אתה כותב {kind} עבור בעל עסק, בקול שלו.

חובה לקרוא לפני שאתה כותב:
{format_ref}
{SKILL_REFS}/privacy.md — הכללים על שמות

⚠️ סדר עדיפויות כשיש התנגשות: הקול > הסודיות > מבנה הפורמט.

⚠️ אסור להזכיר שמות של אנשים, עסקים או מקומות. ציטוטים מותרים בלי מקור.
⚠️ אסור להמציא ציטוטים. רק מה שבניתוח.

⚠️⚠️ **השתמש בשני הקולות שבניתוח, לא רק באחד:**
- **הכאב והציטוט של הלקוח** — זה מה שגורם לקורא לעצור ולהגיד "זה אני"
- **התשובה, האבחון או הסיפור של בעל העסק** — זה מה שהופך את הפוסט
  מתיאור בעיה לעמדה של מי שיודע לפתור אותה

פיסת תוכן שיש בה רק את הכאב מתארת בעיה ונגמרת. פיסה שיש בה רק את
התשובה נשמעת כמו הרצאה. **שתיהן ביחד זה מה שעובד.**

ואם יש בניתוח סיפור אישי של בעל העסק — **שלב אותו.** זה הדבר שהכי
מהיר בונה חיבור, וזה מה שאף אחד אחר לא יכול להעתיק.

הקול של בעל העסק:
{voice}
{examples_block}{layer_line}

{questions_block}

---
הניתוח של הפגישה:

{analysis}
{mirror_block}

---

⚠️ לפני שאתה מחזיר — עבור על סעיף "מה אסור לי" בקובץ הקול, ובדוק את
הפלט שלך מולו שורה אחר שורה. **כולל ההאשטגים והכותרות.**

זה לא סעיף להתרשמות. אם כתוב שם שמילה מסוימת אסורה, היא אסורה גם
בהאשטג, גם בכותרת, וגם כשהיא "מתאימה בול". בבדיקה אמיתית של הסוכן
הזה, מילה אסורה עברה בדיוק ככה — בהאשטג בתחתית הפוסט, אחרי שכל
גוף הטקסט היה נקי.

החזר רק את הפלט הסופי, אחרי שתיקנת.
"""


if __name__ == "__main__":
    test = "--test" in sys.argv
    try:
        run(test_mode=test)
    except Exception as e:
        setup_logging()
        tb = traceback.format_exc()
        log(f"FATAL: {e}")
        log(tb)
        try:
            (LOGS / "last_failure.txt").write_text(
                f"{datetime.now():%Y-%m-%d %H:%M}\n{e}\n\n{tb}", encoding="utf-8")
        except Exception:
            pass
        sys.exit(1)
